# docx_builder.py
"""
Utilidades para construir el entregable final en formato DOCX y subirlo a Google Drive.

Funciones públicas:
- parse_markdown_table_to_docx(doc, table_lines)
- find_image_by_stem(requested_name, image_map)
- save_final_deliverable(...)
"""

import os
from collections import defaultdict
from typing import List, Dict, Any

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from googleapiclient.http import MediaFileUpload

import logging

logger = logging.getLogger(__name__)

# -------------------------------------------------------------------
# Configuración de secciones del DOCX
# -------------------------------------------------------------------

SECTION_MAP = {
    "[SECCION:: TITULO]": ("", 0),  # el título real se arma con el nombre del cliente
    "[SECCION:: EVENTOS_DE_ABUSO]": ("DESCRIPCIÓN DE EVENTOS DE ABUSO", 1),
    "[SECCION:: WITNESS]": ("WITNESS", 1),
    "[SECCION:: EVENTOS_DE_GMC]": ("DESCRIPCIÓN DE GMC", 1),
    "[SECCION:: REFERENCE_LETTERS]": ("REFERENCE LETTERS", 1),
    "[SECCION:: PERMAMENT_BAR]": ("CUESTIONARIO PERMAMENT BAR", 1),
}

ORDERED_SECTION_TAGS = [
    "[SECCION:: TITULO]",
    "[SECCION:: EVENTOS_DE_ABUSO]",
    "[SECCION:: WITNESS]",
    "[SECCION:: EVENTOS_DE_GMC]",
    "[SECCION:: REFERENCE_LETTERS]",
    "[SECCION:: PERMAMENT_BAR]",
]


# -------------------------------------------------------------------
# Tablas Markdown → python-docx
# -------------------------------------------------------------------

def parse_markdown_table_to_docx(doc: Document, table_lines: List[str]) -> None:
    """
    Parsea una lista de líneas que representan una tabla de Markdown y la inserta
    como una tabla nativa de python-docx dentro del documento dado.

    Ejemplo de table_lines:
        [
            "| Col1 | Col2 |",
            "|------|------|",
            "| A    | B    |",
            "| C    | D    |",
        ]
    """
    if not table_lines:
        return

    # Limpiar líneas: remover barras iniciales/finales, trim.
    data_lines = [
        [cell.strip() for cell in line.strip("|").split("|")]
        for line in table_lines
        if line.strip().startswith("|")
    ]

    if len(data_lines) < 2:
        # Mínimo: Header y Separador
        doc.add_paragraph(
            "[ERROR: Datos de tabla incompletos o mal formados.]"
        )
        return

    # La primera línea es la cabecera
    headers = data_lines[0]
    num_cols = len(headers)

    if num_cols == 0:
        doc.add_paragraph("[ERROR: No se detectaron columnas en la tabla.]")
        return

    # Restamos 1 para el separador (línea |:---|:---|)
    num_rows = len(data_lines) - 1

    # Si la tabla tiene solo 2 líneas (Header y Separador), solo incluimos la cabecera.
    if num_rows <= 0:
        num_rows = 1

    table = doc.add_table(rows=num_rows, cols=num_cols)
    table.style = "Table Grid"

    # Rellenar cabecera
    hdr_cells = table.rows[0].cells
    for i, header_text in enumerate(headers):
        header_text = header_text.replace("**", "").strip()
        hdr_cells[i].text = header_text

    # Rellenar filas de datos
    # Empezamos desde la línea 2 (índice 2) ya que 0 = header, 1 = separador
    for r_idx, row_data in enumerate(data_lines[2:]):
        try:
            cells = table.rows[r_idx + 1].cells
            for c_idx, cell_text in enumerate(row_data):
                if c_idx < num_cols:
                    cells[c_idx].text = cell_text
                else:
                    break
        except IndexError:
            doc.add_paragraph(
                f"[ERROR: Fila {r_idx + 1} de tabla falló al insertarse.]"
            )
            break


# -------------------------------------------------------------------
# Búsqueda de imágenes por nombre
# -------------------------------------------------------------------

def find_image_by_stem(requested_name: str, image_map: Dict[str, str]) -> str | None:
    """
    Busca una imagen en el mapa por su nombre exacto o por su 'stem' (nombre sin extensión).

    Args:
        requested_name: Nombre pedido en el marcador, por ejemplo "foto1.jpg"
                         o "foto1" (sin extensión).
        image_map: Dict {nombre_archivo_en_drive: ruta_local}.

    Returns:
        Ruta local de la imagen si se encuentra, o None.
    """
    # Búsqueda directa
    if requested_name in image_map:
        return image_map[requested_name]

    # Búsqueda por 'stem'
    try:
        requested_stem = os.path.splitext(requested_name)[0].lower()
        for real_name, real_path in image_map.items():
            real_stem = os.path.splitext(real_name)[0].lower()
            if requested_stem == real_stem:
                return real_path
    except Exception:
        pass

    return None


# -------------------------------------------------------------------
# Construcción y subida del entregable final
# -------------------------------------------------------------------

def save_final_deliverable(
    drive_service,
    deliverable_text: str,
    client_name: str,
    parent_folder_id: str,
    temp_dir: str,
    abuse_images: List[Dict[str, Any]],
    gmc_images: List[Dict[str, Any]],
) -> List[Dict[str, Any]]:
    """
    Guarda el texto del entregable final en .txt y .docx, sube los archivos a Drive,
    aplica permisos de editor y devuelve metadatos de los archivos subidos.

    Args:
        drive_service: Cliente de Google Drive (build('drive', 'v3', ...)).
        deliverable_text: Texto generado por la IA.
        client_name: Nombre del cliente (se usa en el título y nombres de archivo).
        parent_folder_id: Carpeta de Drive donde se guardarán los archivos.
        temp_dir: Directorio temporal donde se crean los archivos locales.
        abuse_images: Lista de dicts con información de imágenes de ABUSE.
                      Ej: {'name': 'foto1.jpg', 'path': '/tmp/...', 'id': '...', 'mimeType': 'image/jpeg'}
        gmc_images: Lista similar, para imágenes de GMC.

    Returns:
        Lista de dicts con info de archivos subidos:
        [{'id': '...', 'link': 'https://...', 'name': 'DOE_CLIENTE.docx'}, ...]
    """
    logger.info("✍️ Guardando y formateando entregable final...")
    base_filename = f"DOE_{client_name.upper()}"

    # ------------------------------------------------------------------
    # Mapa de imágenes (nombre -> ruta)
    # ------------------------------------------------------------------
    image_map: Dict[str, str] = {}
    for img in abuse_images + gmc_images:
        image_map[img["name"]] = img["path"]
    logger.info("   ... Mapa de %d imágenes creado para inserción.", len(image_map))

    # ------------------------------------------------------------------
    # Archivo .txt (raw)
    # ------------------------------------------------------------------
    txt_filename = f"{base_filename}_RAW_OUTPUT.txt"
    local_txt_path = os.path.join(temp_dir, txt_filename)
    try:
        with open(local_txt_path, "w", encoding="utf-8") as f:
            f.write(deliverable_text)
        logger.info("   ... Archivo .txt (raw) local creado: %s", local_txt_path)
    except Exception as e:
        logger.error("   ❌ Error al crear .txt local: %s", e)

    # ------------------------------------------------------------------
    # Archivo .docx formateado
    # ------------------------------------------------------------------
    docx_filename = f"{base_filename}.docx"
    local_docx_path = os.path.join(temp_dir, docx_filename)

    try:
        doc = Document()

        # 1) Separar contenido por secciones
        section_content: Dict[str, List[str]] = defaultdict(list)
        current_section_tag: str | None = None

        lines = deliverable_text.splitlines()
        for line in lines:
            line = line.strip()
            is_section_tag = False

            for tag in SECTION_MAP.keys():
                if line.startswith(tag):
                    current_section_tag = tag
                    is_section_tag = True
                    break

            if is_section_tag:
                continue

            if current_section_tag:
                section_content[current_section_tag].append(line)

        # 2) Ensamblar documento en orden
        for tag in ORDERED_SECTION_TAGS:
            if tag not in SECTION_MAP:
                continue

            title, level = SECTION_MAP[tag]

            # Título principal
            if tag == "[SECCION:: TITULO]":
                main_title = f"DOE {client_name.upper()}"
                doc.add_heading(main_title, level=0).alignment = WD_ALIGN_PARAGRAPH.CENTER
                continue

            # Encabezado de sección
            heading = doc.add_heading(title, level=level)
            heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

            table_buffer: List[str] = []
            is_in_table = False

            for line in section_content[tag]:
                is_markdown_row = line.startswith("|") and "|" in line[1:]

                # --- Manejo de tablas en PERMAMENT_BAR ---
                if tag == "[SECCION:: PERMAMENT_BAR]" and is_markdown_row:
                    table_buffer.append(line)
                    is_in_table = True
                    continue

                # Cierre de tabla cuando termina el bloque de filas
                if is_in_table and not is_markdown_row and table_buffer:
                    doc.add_paragraph().add_run(
                        "--- Tabla de Entradas y Salidas ---"
                    ).bold = True
                    parse_markdown_table_to_docx(doc, table_buffer)
                    table_buffer = []
                    is_in_table = False

                if is_in_table:
                    # Si seguimos en tabla pero la línea también es markdown, ya la agregamos arriba
                    continue

                # --- Contenido normal de la sección ---
                if line.startswith("EVENTO::"):
                    try:
                        event_title = line.split("::", 1)[1].strip()
                        doc.add_heading(f"Evento: {event_title}", level=2)
                    except Exception:
                        doc.add_heading("Evento (Error de Formato)", level=2)

                elif line.startswith("[IMAGEN::"):
                    try:
                        image_name = line[9:-1].strip()  # [IMAGEN:: xxx ]
                        image_path = find_image_by_stem(image_name, image_map)

                        if image_path:
                            logger.info("   ... Insertando imagen: %s", image_name)
                            doc.add_picture(image_path, width=Inches(5.0))
                            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                        else:
                            p = doc.add_paragraph()
                            run = p.add_run(
                                f"[Error: Imagen '{image_name}' solicitada pero no encontrada.]"
                            )
                            run.italic = True
                            run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)

                    except Exception as e:
                        logger.error("   ❌ Error al procesar/insertar imagen %s: %s", line, e)
                        p = doc.add_paragraph(
                            f"[Error al procesar marcador de imagen: {e}]"
                        )
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

                elif line.startswith("DESCRIPCION::"):
                    try:
                        desc_text = line.split("::", 1)[1].strip()
                        p = doc.add_paragraph()
                        run = p.add_run(desc_text)
                        run.font.name = "Times New Roman"
                        run.font.size = Pt(10)
                        run.italic = True
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    except Exception as e:
                        logger.error("   ❌ Error al procesar descripción de imagen: %s", e)
                        doc.add_paragraph(
                            f"[Error al procesar descripción: {e}]"
                        )

                else:
                    # Párrafo normal
                    p = doc.add_paragraph()
                    run = p.add_run(line)
                    run.font.name = "Times New Roman"
                    if tag == "[SECCION:: PERMAMENT_BAR]":
                        run.font.size = Pt(10)
                    else:
                        run.font.size = Pt(12)
                    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

            # Si la sección terminó todavía dentro de una tabla, cerrarla
            if is_in_table and table_buffer:
                doc.add_paragraph().add_run(
                    "--- Tabla de Entradas y Salidas ---"
                ).bold = True
                parse_markdown_table_to_docx(doc, table_buffer)

        doc.save(local_docx_path)
        logger.info("   ... Archivo .docx FORMATEADO local creado: %s", local_docx_path)

    except Exception as e:
        logger.error("   ❌ Error al crear .docx formateado: %s", e)

    # ------------------------------------------------------------------
    # Subir archivos a Drive
    # ------------------------------------------------------------------
    files_to_upload = [
        (local_txt_path, "text/plain"),
        (
            local_docx_path,
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        ),
    ]

    uploaded_files_info: List[Dict[str, Any]] = []

    for local_path, mime_type in files_to_upload:
        if not os.path.exists(local_path):
            continue

        try:
            media = MediaFileUpload(local_path, mimetype=mime_type)
            file_metadata = {
                "name": os.path.basename(local_path),
                "parents": [parent_folder_id],
            }

            uploaded_file = (
                drive_service.files()
                .create(
                    body=file_metadata,
                    media_body=media,
                    fields="id, webViewLink",
                )
                .execute()
            )

            file_id = uploaded_file.get("id")
            file_link = uploaded_file.get("webViewLink")

            logger.info(
                "   ... Archivo '%s' subido a Drive (ID: %s).",
                os.path.basename(local_path),
                file_id,
            )

            # Aplicar permisos de editor (anyone writer)
            if file_id:
                try:
                    permission_body = {
                        "type": "anyone",
                        "role": "writer",
                    }
                    drive_service.permissions().create(
                        fileId=file_id,
                        body=permission_body,
                        fields="id",
                    ).execute()
                    logger.info(
                        "   ... Permisos de 'editor' (anyone) aplicados a %s.",
                        file_id,
                    )
                except Exception as e:
                    logger.warning(
                        "   ⚠️ Error al aplicar permisos a %s: %s",
                        file_id,
                        e,
                    )

            uploaded_files_info.append(
                {
                    "id": file_id,
                    "link": file_link,
                    "name": os.path.basename(local_path),
                }
            )

        except Exception as e:
            logger.error(
                "   ❌ Error al subir '%s' a Drive: %s",
                os.path.basename(local_path),
                e,
            )

    return uploaded_files_info
